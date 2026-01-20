"""
Document Parser Service
Extracts text from PDF, DOCX, and LaTeX files for ATS analysis.
"""
import logging
import io
import re
from typing import Tuple

logger = logging.getLogger(__name__)


class DocumentParser:
    """Service for extracting text from various document formats"""

    SUPPORTED_FORMATS = {'.pdf', '.docx', '.doc', '.tex', '.latex'}

    def detect_file_type(self, filename: str, content_type: str = None) -> str:
        """Detect file type from filename and content type"""
        filename_lower = filename.lower()

        if filename_lower.endswith('.pdf') or content_type == 'application/pdf':
            return 'pdf'
        elif filename_lower.endswith(('.docx', '.doc')) or content_type in [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword'
        ]:
            return 'docx'
        elif filename_lower.endswith(('.tex', '.latex')):
            return 'latex'
        else:
            return 'unknown'

    async def extract_text(self, file_bytes: bytes, filename: str, content_type: str = None) -> Tuple[str, str]:
        """
        Extract text from document.
        Returns tuple of (extracted_text, file_type)
        """
        file_type = self.detect_file_type(filename, content_type)

        if file_type == 'pdf':
            text = await self.extract_text_from_pdf(file_bytes)
        elif file_type == 'docx':
            text = await self.extract_text_from_docx(file_bytes)
        elif file_type == 'latex':
            text = await self.extract_text_from_latex(file_bytes)
        else:
            raise ValueError(f"Unsupported file format: {filename}")

        return text, file_type

    async def extract_text_from_pdf(self, file_bytes: bytes) -> str:
        """Extract text from PDF file"""
        try:
            import PyPDF2

            pdf_file = io.BytesIO(file_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            text_parts = []
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

            full_text = '\n'.join(text_parts)
            logger.info(f"Extracted {len(full_text)} characters from PDF ({len(pdf_reader.pages)} pages)")
            return full_text

        except ImportError:
            logger.error("PyPDF2 not installed. Install with: pip install PyPDF2")
            raise ValueError("PDF parsing not available. Server configuration error.")
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise ValueError(f"Failed to parse PDF: {str(e)}")

    async def extract_text_from_docx(self, file_bytes: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            from docx import Document

            docx_file = io.BytesIO(file_bytes)
            doc = Document(docx_file)

            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))

            full_text = '\n'.join(text_parts)
            logger.info(f"Extracted {len(full_text)} characters from DOCX")
            return full_text

        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            raise ValueError("DOCX parsing not available. Server configuration error.")
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise ValueError(f"Failed to parse DOCX: {str(e)}")

    async def extract_text_from_latex(self, file_bytes: bytes) -> str:
        """Extract readable text from LaTeX file, removing commands"""
        try:
            latex_content = file_bytes.decode('utf-8', errors='ignore')

            # Remove comments
            text = re.sub(r'%.*$', '', latex_content, flags=re.MULTILINE)

            # Remove common LaTeX commands while preserving content
            # Remove documentclass, usepackage, etc.
            text = re.sub(r'\\documentclass\[?[^\]]*\]?\{[^}]*\}', '', text)
            text = re.sub(r'\\usepackage\[?[^\]]*\]?\{[^}]*\}', '', text)
            text = re.sub(r'\\begin\{document\}', '', text)
            text = re.sub(r'\\end\{document\}', '', text)

            # Extract content from common commands
            text = re.sub(r'\\section\*?\{([^}]*)\}', r'\n\1\n', text)
            text = re.sub(r'\\subsection\*?\{([^}]*)\}', r'\n\1\n', text)
            text = re.sub(r'\\textbf\{([^}]*)\}', r'\1', text)
            text = re.sub(r'\\textit\{([^}]*)\}', r'\1', text)
            text = re.sub(r'\\emph\{([^}]*)\}', r'\1', text)
            text = re.sub(r'\\href\{[^}]*\}\{([^}]*)\}', r'\1', text)
            text = re.sub(r'\\url\{([^}]*)\}', r'\1', text)

            # Remove remaining LaTeX commands
            text = re.sub(r'\\[a-zA-Z]+\*?\{[^}]*\}', '', text)
            text = re.sub(r'\\[a-zA-Z]+\[[^\]]*\]', '', text)
            text = re.sub(r'\\[a-zA-Z]+', '', text)

            # Remove braces, special chars
            text = re.sub(r'[{}]', '', text)
            text = re.sub(r'\\[&%$#_]', '', text)

            # Clean up whitespace
            text = re.sub(r'\n\s*\n', '\n\n', text)
            text = re.sub(r'  +', ' ', text)
            text = text.strip()

            logger.info(f"Extracted {len(text)} characters from LaTeX")
            return text

        except Exception as e:
            logger.error(f"Error extracting text from LaTeX: {str(e)}")
            raise ValueError(f"Failed to parse LaTeX: {str(e)}")


# Global instance
document_parser = DocumentParser()
